"""
build_metodichka_docx.py
========================
Генерирует docs/metodichka/Методичка_подробная.docx —
полную методичку с разбором кода, готовую к редактированию и сдаче.

Запуск:
    python3 scripts/build_metodichka_docx.py
"""

from __future__ import annotations

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ────────────────────────────────────────────────────────────
# Пути
# ────────────────────────────────────────────────────────────

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_ROOT = os.path.dirname(SCRIPT_DIR)
OUTPUT_PATH = os.path.join(PROJECT_ROOT, "docs", "metodichka", "Методичка_подробная.docx")


# ────────────────────────────────────────────────────────────
# Цвета (RGB)
# ────────────────────────────────────────────────────────────

COLOR_DARK      = RGBColor(0x1A, 0x1A, 0x2E)   # почти чёрный — заголовки
COLOR_ACCENT    = RGBColor(0x16, 0x21, 0x3E)   # тёмно-синий — H2
COLOR_BLUE      = RGBColor(0x0F, 0x3C, 0x78)   # синий — H3
COLOR_CODE_BG   = RGBColor(0xF0, 0xF4, 0xF8)   # светло-серый — фон кода
COLOR_CODE_FG   = RGBColor(0x1E, 0x3A, 0x5F)   # тёмно-синий — текст кода
COLOR_NOTE_BG   = RGBColor(0xFF, 0xF9, 0xE6)   # светло-жёлтый — примечания
COLOR_NOTE_FG   = RGBColor(0x7A, 0x5C, 0x00)   # золотистый — текст примечаний
COLOR_TABLE_HDR = RGBColor(0x16, 0x21, 0x3E)   # заголовки таблиц
COLOR_GRAY_ALT  = RGBColor(0xF5, 0xF7, 0xFA)   # чередующиеся строки таблицы


# ────────────────────────────────────────────────────────────
# Вспомогательные функции
# ────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str) -> None:
    """Устанавливает цвет фона ячейки таблицы."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_border(cell, **kwargs) -> None:
    """Устанавливает рамку ячейки. kwargs: top, bottom, left, right."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, val in kwargs.items():
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), val.get("val", "single"))
        border.set(qn("w:sz"), str(val.get("sz", 4)))
        border.set(qn("w:color"), val.get("color", "auto"))
        tcBorders.append(border)
    tcPr.append(tcBorders)


def add_run_bold(para, text: str, size: int = 11, color: RGBColor | None = None) -> None:
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color


def add_run_normal(para, text: str, size: int = 11) -> None:
    run = para.add_run(text)
    run.font.size = Pt(size)


def add_horizontal_rule(doc: Document) -> None:
    """Добавляет горизонтальную линию-разделитель."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(4)
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "C0C8D8")
    pBdr.append(bottom)
    pPr.append(pBdr)


# ────────────────────────────────────────────────────────────
# Стили параграфов
# ────────────────────────────────────────────────────────────

def h1(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(24)
    para.paragraph_format.space_after = Pt(8)
    para.paragraph_format.keep_with_next = True
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = COLOR_DARK
    add_horizontal_rule(doc)


def h2(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(16)
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.keep_with_next = True
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = COLOR_ACCENT


def h3(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(4)
    para.paragraph_format.keep_with_next = True
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = COLOR_BLUE


def body(doc: Document, text: str, bold_parts: list[str] | None = None) -> None:
    """Обычный параграф. bold_parts — список подстрок для выделения жирным."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(3)
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.first_line_indent = Cm(0.75)

    if not bold_parts:
        run = para.add_run(text)
        run.font.size = Pt(11)
    else:
        remaining = text
        for bp in bold_parts:
            idx = remaining.find(bp)
            if idx == -1:
                continue
            if idx > 0:
                r = para.add_run(remaining[:idx])
                r.font.size = Pt(11)
            r2 = para.add_run(bp)
            r2.bold = True
            r2.font.size = Pt(11)
            remaining = remaining[idx + len(bp):]
        if remaining:
            r = para.add_run(remaining)
            r.font.size = Pt(11)


def body_plain(doc: Document, text: str) -> None:
    """Параграф без отступа первой строки."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(3)
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run(text)
    run.font.size = Pt(11)


def bullet(doc: Document, text: str, level: int = 0, bold_part: str = "") -> None:
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.left_indent = Cm(0.75 + level * 0.5)
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)
    if bold_part and bold_part in text:
        idx = text.find(bold_part)
        if idx > 0:
            r = para.add_run(text[:idx])
            r.font.size = Pt(11)
        r2 = para.add_run(bold_part)
        r2.bold = True
        r2.font.size = Pt(11)
        rest = text[idx + len(bold_part):]
        if rest:
            r3 = para.add_run(rest)
            r3.font.size = Pt(11)
    else:
        run = para.add_run(text)
        run.font.size = Pt(11)


def numbered(doc: Document, text: str, bold_part: str = "") -> None:
    para = doc.add_paragraph(style="List Number")
    para.paragraph_format.left_indent = Cm(0.75)
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)
    if bold_part and bold_part in text:
        idx = text.find(bold_part)
        if idx > 0:
            r = para.add_run(text[:idx])
            r.font.size = Pt(11)
        r2 = para.add_run(bold_part)
        r2.bold = True
        r2.font.size = Pt(11)
        rest = text[idx + len(bold_part):]
        if rest:
            r3 = para.add_run(rest)
            r3.font.size = Pt(11)
    else:
        run = para.add_run(text)
        run.font.size = Pt(11)


def code_block(doc: Document, lines: list[str], caption: str = "") -> None:
    """Блок кода с монопространным шрифтом и серым фоном."""
    if caption:
        cp = doc.add_paragraph()
        cp.paragraph_format.space_before = Pt(10)
        cp.paragraph_format.space_after = Pt(0)
        r = cp.add_run(f"  {caption}")
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x55, 0x66, 0x77)

    for line in lines:
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.left_indent = Cm(0.5)
        para.paragraph_format.right_indent = Cm(0.5)
        # Фон через shading
        pPr = para._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "EEF2F8")
        pPr.append(shd)
        run = para.add_run(line if line else " ")
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        run.font.color.rgb = COLOR_CODE_FG

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def note_box(doc: Document, text: str, prefix: str = "📌 Для преподавателя") -> None:
    """Блок-примечание с жёлтым фоном."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after = Pt(8)
    para.paragraph_format.left_indent = Cm(0.5)
    para.paragraph_format.right_indent = Cm(0.5)
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "FFF8DC")
    pPr.append(shd)
    r1 = para.add_run(f"{prefix}:  ")
    r1.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = RGBColor(0x7A, 0x5C, 0x00)
    r2 = para.add_run(text)
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(0x4A, 0x3C, 0x00)


def screenshot_placeholder(doc: Document, description: str) -> None:
    """Место для скриншота."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(10)
    para.paragraph_format.space_after = Pt(10)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "E8F4FF")
    pPr.append(shd)
    r1 = para.add_run("📸  МЕСТО ДЛЯ СКРИНШОТА\n")
    r1.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)
    r2 = para.add_run(description)
    r2.italic = True
    r2.font.size = Pt(9)
    r2.font.color.rgb = RGBColor(0x1A, 0x5C, 0xAA)


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    col_widths: list[float] | None = None,
) -> None:
    """Добавляет оформленную таблицу."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"

    # Заголовок
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, "162130")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Данные
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        bg = "F5F7FA" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            set_cell_bg(cell, bg)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            para = cell.paragraphs[0]
            run = para.add_run(cell_text)
            run.font.size = Pt(10)

    # Ширины столбцов
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)


# ────────────────────────────────────────────────────────────
# Титульная страница
# ────────────────────────────────────────────────────────────

def build_title_page(doc: Document) -> None:
    doc.add_paragraph().paragraph_format.space_after = Pt(60)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run("STEM-ПРОЕКТ")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x55, 0x66, 0x88)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = t2.add_run("3D-визуализация математических поверхностей\nи поиск оптимального пути в Blender")
    run2.bold = True
    run2.font.size = Pt(22)
    run2.font.color.rgb = COLOR_DARK

    doc.add_paragraph().paragraph_format.space_after = Pt(20)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Подробная методическая разработка")
    r.italic = True
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0x44, 0x55, 0x77)

    add_horizontal_rule(doc)
    doc.add_paragraph().paragraph_format.space_after = Pt(40)

    for line, size, bold in [
        ("Предмет: математика + информатика + 3D-технологии", 12, False),
        ("Класс: 10–11 / первый курс", 12, False),
        ("Инструменты: Blender 4.x/5.x, Python 3.10+", 12, False),
        ("Ориентировочное время: 6–8 учебных часов", 12, False),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        r.bold = bold
        r.font.size = Pt(size)

    doc.add_paragraph().paragraph_format.space_after = Pt(60)
    doc.add_page_break()


# ────────────────────────────────────────────────────────────
# Оглавление (ручное)
# ────────────────────────────────────────────────────────────

def build_toc(doc: Document) -> None:
    h1(doc, "Содержание")

    toc_items = [
        ("Глава 0.", "Зачем это всё? Мотивация и цели"),
        ("Глава 1.", "Математика поверхностей: z = f(x, y)"),
        ("Глава 2.", "Как Blender строит 3D из Python"),
        ("Глава 3.", "Разбор function_library.py"),
        ("Глава 4.", "Разбор visualize_function.py"),
        ("Глава 5.", "Geometry Nodes: интерактивная математика"),
        ("Глава 6.", "Поиск пути на поверхности (A* / Dijkstra)"),
        ("Глава 7.", "Эксперименты и таблицы результатов"),
        ("Глава 8.", "Задания для самостоятельной работы"),
        ("Глава 9.", "Оформление результата: отчёт и презентация"),
        ("Глава 10.", "Лабиринт: один и тот же движок на 2D"),
        ("Приложение А.", "Типичные ошибки и их решения"),
        ("Приложение Б.", "Глоссарий терминов"),
    ]

    for num, title in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        r1 = p.add_run(f"{num}  ")
        r1.bold = True
        r1.font.size = Pt(11)
        r1.font.color.rgb = COLOR_ACCENT
        r2 = p.add_run(title)
        r2.font.size = Pt(11)

    doc.add_page_break()


# ────────────────────────────────────────────────────────────
# Глава 0 — Мотивация
# ────────────────────────────────────────────────────────────

def build_ch0(doc: Document) -> None:
    h1(doc, "Глава 0. Зачем это всё?")

    h2(doc, "0.1 Проблема абстрактной математики")
    body(doc,
        "Возьмём формулу из учебника: z = sin(x) · cos(y). Что она описывает? "
        "На бумаге — просто символы. Многие ученики механически подставляют цифры, "
        "не чувствуя формы функции.",
        bold_parts=["z = sin(x) · cos(y)"])
    body(doc,
        "А ведь эта формула описывает волновую поверхность — рябь на воде, "
        "распространяющуюся сразу по двум направлениям одновременно.",
        bold_parts=["волновую поверхность"])

    h2(doc, "0.2 Наш ответ")
    body(doc, "Мы напишем программу на Python, которая:")
    numbered(doc, "Вычисляет значения функции в тысячах точек")
    numbered(doc, "Строит трёхмерный объект в Blender из этих точек")
    numbered(doc, "Красит его по высоте — чтобы «горы» и «долины» были видны сразу")
    numbered(doc, "Находит оптимальный маршрут по этой поверхности (задача дрона)")

    screenshot_placeholder(doc,
        "FunctionVisualizer.blend открыт в Blender — видна цветная волновая поверхность "
        "слева и GN-поверхность справа. Оба объекта выделены, виден 3D Viewport.")

    h2(doc, "0.3 Образовательные результаты")
    body(doc, "После работы с проектом ученик умеет:")
    bullet(doc, "Объяснить, что такое z = f(x, y) и какие виды поверхностей бывают",
           bold_part="z = f(x, y)")
    bullet(doc, "Написать Python-скрипт, который строит 3D-объект из математической формулы")
    bullet(doc, "Понять, как работает Geometry Nodes — «умный» модификатор Blender")
    bullet(doc, "Объяснить алгоритмы поиска пути (Dijkstra, A*) и применить их к реальной задаче")
    bullet(doc, "Провести вычислительный эксперимент и зафиксировать результат в таблице")

    note_box(doc,
        "Проект рассчитан на 6–8 учебных часов. Главы 0–4 — базовый модуль (4 ч). "
        "Главы 5–6 — расширенный модуль (2–3 ч). Задания главы 8 — домашняя работа.")

    doc.add_page_break()


# ────────────────────────────────────────────────────────────
# Глава 1 — Математика
# ────────────────────────────────────────────────────────────

def build_ch1(doc: Document) -> None:
    h1(doc, "Глава 1. Математика поверхностей")

    h2(doc, "1.1 Что такое z = f(x, y)?")
    body(doc,
        "Обычная функция y = f(x) задаёт кривую на плоскости. "
        "Функция двух переменных z = f(x, y) задаёт поверхность в трёхмерном пространстве. "
        "Для каждой точки (x, y) на плоскости существует ровно одна высота z.",
        bold_parts=["y = f(x)", "z = f(x, y)"])
    body(doc,
        "Пример: z = x² + y². При x=1, y=1 → z=2. При x=0, y=0 → z=0. "
        "Это параболоид — форма параболы, вращённой вокруг вертикальной оси.",
        bold_parts=["z = x² + y²", "параболоид"])

    h2(doc, "1.2 Пять функций проекта")

    add_table(doc,
        headers=["Имя", "Формула", "Форма", "Где встречается"],
        rows=[
            ["paraboloid", "z = A·(x²+y²)",           "«Миска», воронка",       "Параболические антенны, линзы, горки"],
            ["saddle",     "z = A·(x²−y²)",           "«Седло», «чипс»",        "Поверхности минимальной энергии"],
            ["wave",       "z = A·sin(k·x)·cos(k·y)", "«Морская рябь»",         "Волны, звуковые и световые поля"],
            ["ripple",     "z = A·sin(k·√(x²+y²))",  "«Круги на воде»",        "Распространение волн от точечного источника"],
            ["gaussian",   "z = A·exp(−(x²+y²)/σ²)", "«Купол», «горка»",       "Нормальное распределение, тепловые пятна"],
        ],
        col_widths=[2.5, 4.5, 3.5, 5.0],
    )

    h2(doc, "1.3 Параметры и их смысл")
    body(doc, "Каждая функция управляется тремя параметрами:")
    bullet(doc, "A (amplitude) — амплитуда, «высота» формы. Чем больше A, тем выше «горы».",
           bold_part="A (amplitude)")
    bullet(doc, "k (frequency) — частота. Чем больше k, тем «гуще» волны на том же участке.",
           bold_part="k (frequency)")
    bullet(doc, "σ (sigma) — ширина гауссова купола. Большое σ → широкая пологая горка.",
           bold_part="σ (sigma)")

    screenshot_placeholder(doc,
        "Рендеры пяти функций рядом: paraboloid_default.png, wave_A1_k1.png, "
        "saddle_default.png, gaussian_sigma2.png, ripple_k1.png — из папки assets/renders/.")

    h2(doc, "1.4 Задание для самостоятельной работы")
    note_box(doc,
        "Нарисуй от руки, как, по-твоему, выглядит каждая из 5 функций. "
        "Затем запусти проект и сравни с результатом. Где ты угадал?",
        prefix="✏️ Задание 1")

    doc.add_page_break()


# ────────────────────────────────────────────────────────────
# Глава 2 — Как Blender строит 3D
# ────────────────────────────────────────────────────────────

def build_ch2(doc: Document) -> None:
    h1(doc, "Глава 2. Как Blender строит 3D из Python")

    h2(doc, "2.1 Что такое mesh?")
    body(doc,
        "Трёхмерный объект в Blender состоит из трёх базовых элементов:")
    bullet(doc, "Вершины (vertices) — точки в пространстве с координатами (x, y, z)",
           bold_part="Вершины (vertices)")
    bullet(doc, "Рёбра (edges) — линии между вершинами",
           bold_part="Рёбра (edges)")
    bullet(doc, "Грани (faces) — плоские многоугольники, ограниченные рёбрами",
           bold_part="Грани (faces)")
    body(doc,
        "Чтобы построить поверхность z = f(x, y), мы создаём сетку точек на плоскости, "
        "для каждой вычисляем z, получаем вершины (x, y, z), соединяем их в квадратные грани "
        "и передаём в Blender.")

    screenshot_placeholder(doc,
        "Blender в режиме Edit Mode — виден wireframe сетки поверхности. "
        "Переход: Tab → Edit Mode, затем Overlays → Wire.")

    h2(doc, "2.2 Минимальный пример: куб из Python")
    body(doc,
        "Прежде чем разбирать большой скрипт, посмотрим на минимальный рабочий пример — "
        "построение куба из 8 вершин и 6 граней:")

    code_block(doc, [
        "import bpy",
        "",
        "# 8 вершин куба",
        "vertices = [",
        "    (-1, -1, -1), (1, -1, -1), (1, 1, -1), (-1, 1, -1),  # нижняя грань",
        "    (-1, -1,  1), (1, -1,  1), (1, 1,  1), (-1, 1,  1),  # верхняя грань",
        "]",
        "",
        "# 6 граней — каждая задаётся 4 индексами вершин",
        "faces = [",
        "    (0,1,2,3), (4,5,6,7),  # низ и верх",
        "    (0,1,5,4), (2,3,7,6),  # перед и зад",
        "    (1,2,6,5), (0,3,7,4),  # право и лево",
        "]",
        "",
        "mesh = bpy.data.meshes.new('MyCube_mesh')",
        "mesh.from_pydata(vertices, [], faces)  # [] = нет отдельных рёбер",
        "mesh.update()",
        "",
        "obj = bpy.data.objects.new('MyCube', mesh)",
        "bpy.context.collection.objects.link(obj)  # добавляем в сцену",
    ], caption="Минимальный пример: куб из 8 вершин")

    h3(doc, "Разбор ключевых строк")
    add_table(doc,
        headers=["Строка кода", "Что делает"],
        rows=[
            ["bpy.data.meshes.new(...)", "Создаёт пустой mesh в памяти Blender"],
            ["mesh.from_pydata(v, [], f)", "Заполняет mesh вершинами и гранями"],
            ["mesh.update()", "Пересчитывает геометрию (нормали и т.д.)"],
            ["bpy.data.objects.new(...)", "Создаёт объект-«обёртку» для mesh"],
            ["collection.objects.link(obj)", "Добавляет объект в текущую сцену"],
        ],
        col_widths=[6.0, 10.0],
    )

    doc.add_page_break()


# ────────────────────────────────────────────────────────────
# Глава 3 — function_library.py
# ────────────────────────────────────────────────────────────

def build_ch3(doc: Document) -> None:
    h1(doc, "Глава 3. Разбор function_library.py")

    h2(doc, "3.1 Зачем отдельный модуль?")
    body(doc,
        "Без function_library.py каждый скрипт содержал бы свою копию формул, "
        "валидации и параметров. Это называется дублирование кода — плохая практика: "
        "при изменении формулы нужно менять её везде и легко что-то пропустить.",
        bold_parts=["дублирование кода"])
    body(doc,
        "Принцип единственного источника правды (Single Source of Truth): "
        "все формулы, все параметры по умолчанию, вся валидация — только здесь.",
        bold_parts=["Принцип единственного источника правды"])

    h2(doc, "3.2 SurfaceConfig — хранилище параметров")
    code_block(doc, [
        "@dataclass(frozen=True, slots=True)",
        "class SurfaceConfig:",
        "    function:   str   = 'wave'   # имя функции",
        "    resolution: int   = 60       # 60×60 точек сетки",
        "    x_min:      float = -5.0     # область по X",
        "    x_max:      float =  5.0",
        "    y_min:      float = -5.0     # область по Y",
        "    y_max:      float =  5.0",
        "    amplitude:  float =  1.0     # A — высота",
        "    frequency:  float =  1.0     # k — частота волн",
        "    sigma:      float =  2.0     # σ — ширина купола",
    ], caption="scripts/function_library.py — класс SurfaceConfig")

    h3(doc, "Разбор деталей")
    body(doc,
        "@dataclass — декоратор Python, который автоматически создаёт конструктор __init__, "
        "метод __repr__ и другие. Нам не нужно писать их вручную.",
        bold_parts=["@dataclass"])
    body(doc,
        "frozen=True делает объект неизменяемым. После создания изменить параметры нельзя — "
        "Python выбросит FrozenInstanceError. Это защита: параметры не должны меняться "
        "во время рендера.",
        bold_parts=["frozen=True", "FrozenInstanceError"])
    body(doc,
        "slots=True — оптимизация памяти. Вместо словаря Python использует "
        "__slots__, что быстрее и занимает меньше RAM.",
        bold_parts=["slots=True", "__slots__"])

    h2(doc, "3.3 Реестр функций")
    code_block(doc, [
        "FUNCTION_REGISTRY: dict[str, SurfaceFormula] = {",
        "    'paraboloid': _paraboloid,",
        "    'saddle':     _saddle,",
        "    'wave':       _wave,",
        "    'ripple':     _ripple,",
        "    'gaussian':   _gaussian,",
        "    'custom':     _custom,",
        "}",
    ], caption="Реестр функций — словарь «имя → функция»")

    body(doc,
        "Словарь, где ключ — строка с именем, а значение — сама функция. "
        "В Python функции — объекты первого класса: их можно передавать и хранить в переменных.")
    body(doc,
        "Когда пользователь пишет --function paraboloid, достаточно сделать "
        "FUNCTION_REGISTRY['paraboloid'] — и получить нужную функцию. "
        "Не нужно длинных цепочек if/elif.",
        bold_parts=["--function paraboloid", "FUNCTION_REGISTRY['paraboloid']"])

    h2(doc, "3.4 Математические функции")
    code_block(doc, [
        "def _wave(x: float, y: float, config: SurfaceConfig) -> float:",
        "    k = config.frequency",
        "    return config.amplitude * math.sin(k * x) * math.cos(k * y)",
        "",
        "def _gaussian(x: float, y: float, config: SurfaceConfig) -> float:",
        "    sigma_sq = config.sigma ** 2",
        "    return config.amplitude * math.exp(-(x**2 + y**2) / sigma_sq)",
    ], caption="Примеры математических функций")

    body(doc,
        "Для wave: sin(k·x)·cos(k·y) создаёт узор из стоячих волн. "
        "При k=1 — одна волна на участок [-5, 5]. При k=3 — три волны.",
        bold_parts=["sin(k·x)·cos(k·y)"])
    body(doc,
        "Для gaussian: это двумерное нормальное распределение без нормировки. "
        "sigma_sq = σ² — дисперсия. Чем больше σ, тем шире «купол».",
        bold_parts=["sigma_sq = σ²"])

    h2(doc, "3.5 Генерация геометрии — ключевая функция")
    code_block(doc, [
        "def generate_surface_geometry(config: SurfaceConfig):",
        "    step_x = config.x_span / config.resolution   # шаг по X",
        "    step_y = config.y_span / config.resolution   # шаг по Y",
        "",
        "    for j in range(config.resolution + 1):       # строки сетки",
        "        y = config.y_min + j * step_y",
        "        for i in range(config.resolution + 1):   # столбцы сетки",
        "            x = config.x_min + i * step_x",
        "            z = formula(x, y, config)            # ВЫЧИСЛЯЕМ z",
        "            vertices.append((x, y, z))           # добавляем вершину",
        "",
        "    for j in range(config.resolution):           # строим грани",
        "        for i in range(config.resolution):",
        "            v0 = j * row_len + i",
        "            faces.append((v0, v0+1, v0+row_len+1, v0+row_len))",
    ], caption="Генерация вершин и граней поверхности")

    body(doc,
        "При resolution=60: шаг step_x = 10/60 ≈ 0.167. "
        "Двойной цикл создаёт 61×61 = 3721 вершину и 60×60 = 3600 граней.")

    h2(doc, "3.6 Валидация параметров")
    code_block(doc, [
        "def validate_surface_config(config: SurfaceConfig) -> None:",
        "    get_surface_function(config.function)    # функция существует?",
        "    if config.resolution < 2:",
        "        raise ValueError('resolution должен быть >= 2')",
        "    if config.x_max <= config.x_min:",
        "        raise ValueError('x_max должен быть больше x_min')",
        "    for name, value in finite_params.items():",
        "        if not math.isfinite(value):         # нет inf и nan",
        "            raise ValueError(f'{name} должен быть конечным')",
    ], caption="Валидация — проверка всех параметров до начала работы")

    body(doc,
        "math.isfinite(value) возвращает False, если число равно inf, -inf или nan. "
        "Это важно: если пользователь задаст --amplitude inf, мы выдадим понятную ошибку "
        "вместо непредсказуемого падения посередине рендера.",
        bold_parts=["math.isfinite(value)"])

    doc.add_page_break()


# ────────────────────────────────────────────────────────────
# Глава 4 — visualize_function.py
# ────────────────────────────────────────────────────────────

def build_ch4(doc: Document) -> None:
    h1(doc, "Глава 4. Разбор visualize_function.py")

    h2(doc, "4.1 Структура скрипта")
    code_block(doc, [
        "# 1. Импорт из function_library",
        "from function_library import SurfaceConfig, generate_surface_geometry, ...",
        "",
        "# 2. Попытка импортировать bpy — Blender API",
        "try:",
        "    import bpy",
        "    HAS_BPY = True",
        "except ImportError:",
        "    HAS_BPY = False",
        "    print('[INFO] bpy не найден — режим preview')",
        "",
        "# 3. Функции: создание объекта, материал, камера, рендер",
        "# 4. Точка входа: main()",
    ], caption="Общая структура скрипта")

    body(doc,
        "Конструкция try/except на импорте bpy — это намеренная защита. "
        "bpy существует только внутри Blender. Если запустить скрипт в обычном Python, "
        "импорт упадёт. Мы перехватываем это и продолжаем в режиме preview.",
        bold_parts=["try/except", "bpy", "режиме preview"])

    h2(doc, "4.2 Создание объекта в Blender")
    code_block(doc, [
        "def create_surface_object(vertices, faces, name):",
        "    remove_existing_object(name)       # удаляем старый объект",
        "",
        "    mesh = bpy.data.meshes.new(f'{name}_mesh')",
        "    mesh.from_pydata(vertices, [], faces)",
        "    mesh.update()",
        "",
        "    obj = bpy.data.objects.new(name, mesh)",
        "    bpy.context.collection.objects.link(obj)  # в сцену",
        "    bpy.context.view_layer.objects.active = obj",
        "    obj.select_set(True)",
        "    return obj",
    ], caption="Создание 3D-объекта из вершин и граней")

    body(doc,
        "remove_existing_object удаляет старый объект с таким именем перед созданием нового. "
        "Без этого при многократном запуске скрипт создаёт дублирующиеся объекты в сцене.",
        bold_parts=["remove_existing_object"])

    h2(doc, "4.3 Цветовой материал (colormap)")
    body(doc,
        "Материал — это правило взаимодействия объекта со светом. В Blender материалы строятся "
        "из нодов (узлов) в Shader Editor. Мы строим цепочку:")
    body(doc,
        "Geometry → Separate XYZ (берём Z) → Map Range (z_min..z_max → 0..1) → "
        "Color Ramp (0=синий, 0.5=зелёный, 1=красный) → Principled BSDF → Output",
        bold_parts=["Geometry", "Separate XYZ", "Map Range", "Color Ramp", "Principled BSDF"])

    code_block(doc, [
        "# ColorRamp: три ключевых цвета по высоте",
        "color_ramp.color_ramp.elements[0].color = (0.05, 0.20, 0.85, 1.0)  # синий (низ)",
        "color_ramp.color_ramp.elements[1].color = (0.90, 0.10, 0.10, 1.0)  # красный (верх)",
        "mid = color_ramp.color_ramp.elements.new(0.5)",
        "mid.color                                = (0.10, 0.80, 0.20, 1.0)  # зелёный",
        "",
        "# Map Range: нормализуем z в диапазон [0, 1]",
        "map_range.inputs['From Min'].default_value = z_min",
        "map_range.inputs['From Max'].default_value = z_max",
    ], caption="Настройка цветовой карты по высоте Z")

    screenshot_placeholder(doc,
        "Shader Editor с нод-деревом материала: "
        "Geometry → SeparateXYZ → MapRange → ColorRamp → BSDF → Output.")

    h2(doc, "4.4 Запуск скрипта")
    code_block(doc, [
        "# Внутри Blender: Scripting → Open → Run Script (▶)",
        "",
        "# CLI — фоновый рендер PNG:",
        "blender --background --python scripts/visualize_function.py -- \\",
        "    --function wave --resolution 100 \\",
        "    --amplitude 1.5 --frequency 2 \\",
        "    --output assets/renders/wave_lesson.png",
    ], caption="Два способа запуска скрипта")

    screenshot_placeholder(doc,
        "Blender GUI → вкладка Scripting → скрипт открыт в Text Editor → кнопка Run Script.")
    screenshot_placeholder(doc,
        "Результат в 3D Viewport — цветная волновая поверхность. "
        "Синяя снизу, зелёная в середине, красная сверху.")

    doc.add_page_break()


# ────────────────────────────────────────────────────────────
# Глава 5 — Geometry Nodes
# ────────────────────────────────────────────────────────────

def build_ch5(doc: Document) -> None:
    h1(doc, "Глава 5. Geometry Nodes: интерактивная математика")

    h2(doc, "5.1 Что такое Geometry Nodes?")
    body(doc,
        "Geometry Nodes (GN) — система в Blender, где вместо написания кода "
        "ты визуально соединяешь блоки (ноды), описывающие вычисления над геометрией. "
        "Наш GN-модификатор читает позицию каждой вершины сетки и смещает её по Z "
        "по формуле волны.",
        bold_parts=["Geometry Nodes (GN)"])
    body(doc,
        "Главное преимущество: параметры Amplitude и Frequency можно менять "
        "в реальном времени ползунками, без перезапуска скрипта.",
        bold_parts=["Amplitude", "Frequency"])

    screenshot_placeholder(doc,
        "Объект MathSurface_GN выбран → Properties → вкладка Modifier (синий гаечный ключ) → "
        "модификатор GeoNodes_Surface с полями Amplitude и Frequency.")

    h2(doc, "5.2 Нод-схема волновой поверхности")
    body(doc, "Поток данных в нод-группе:")
    bullet(doc, "GroupInput.Geometry → SetPosition.Geometry → GroupOutput")
    bullet(doc, "InputPosition → SeparateXYZ → X и Y")
    bullet(doc, "X × Frequency → SIN → ")
    bullet(doc, "Y × Frequency → COS → ")
    bullet(doc, "SIN × COS × Amplitude → Z для CombineXYZ")
    bullet(doc, "CombineXYZ(X, Y, Z_волна) → SetPosition.Position")

    code_block(doc, [
        "# Ключевые ноды и их роль:",
        "GeometryNodeInputPosition   — позиция каждой вершины (x, y, z)",
        "ShaderNodeSeparateXYZ       — разделяет вектор на X, Y, Z",
        "ShaderNodeMath (MULTIPLY)   — умножение: k * x, k * y",
        "ShaderNodeMath (SINE)       — вычисляет sin(k*x)",
        "ShaderNodeMath (COSINE)     — вычисляет cos(k*y)",
        "ShaderNodeCombineXYZ        — собирает (X, Y, Z_new) в вектор",
        "GeometryNodeSetPosition     — смещает вершины на новые позиции",
    ], caption="Ноды нод-группы MathSurface_GN_Group")

    screenshot_placeholder(doc,
        "Geometry Node Editor с нод-деревом: "
        "GroupInput → вычисления → SetPosition → GroupOutput. "
        "Виден полный граф с подписями нодов.")

    h2(doc, "5.3 Пресеты параметров")
    body(doc, "В FunctionVisualizer.blend заданы три пресета (хранятся как Custom Properties):")
    add_table(doc,
        headers=["Пресет", "Amplitude", "Frequency", "Назначение"],
        rows=[
            ["lesson_default", "1.0", "1.0", "Стандартный вид для объяснения"],
            ["gentle_wave",    "0.5", "0.7", "Пологая мягкая волна"],
            ["dynamic_wave",   "2.0", "2.5", "Динамичный рельеф"],
        ],
        col_widths=[4.0, 2.5, 2.5, 7.0],
    )
    note_box(doc,
        "Пресеты видны в Properties → Object Properties → Custom Properties. "
        "Чтобы применить пресет: вручную установить значения в модификаторе GeoNodes_Surface.",
        prefix="💡 Совет")

    doc.add_page_break()


# ────────────────────────────────────────────────────────────
# Глава 6 — Pathfinding
# ────────────────────────────────────────────────────────────

def build_ch6(doc: Document) -> None:
    h1(doc, "Глава 6. Поиск пути на поверхности (A* / Dijkstra)")

    h2(doc, "6.1 Постановка задачи")
    body(doc,
        "Дрон летит над горным рельефом, описанным функцией z = f(x, y). "
        "Нужно найти маршрут из точки A в точку B, учитывая:",
        bold_parts=["z = f(x, y)"])
    bullet(doc, "Длину пути в 3D-пространстве")
    bullet(doc, "Штраф за крутые склоны (опасно для дрона)")
    bullet(doc, "Запретные зоны (препятствия)")
    body(doc,
        "Это задача поиска кратчайшего пути на взвешенном графе — одна из фундаментальных "
        "задач информатики, которая используется в навигаторах, играх и робототехнике.",
        bold_parts=["поиска кратчайшего пути на взвешенном графе"])

    h2(doc, "6.2 Граф поверхности (terrain_graph.py)")

    h3(doc, "Дискретизация")
    body(doc,
        "Непрерывную поверхность разбиваем на дискретную сетку. "
        "При resolution=80 получаем 81×81 = 6561 узлов. "
        "Каждый узел — GridNode = (i, j), пара целых индексов.",
        bold_parts=["GridNode = (i, j)"])

    h3(doc, "Связность")
    body(doc,
        "Используем 8-связность: каждый узел соединён с 8 соседями "
        "(4 по вертикали/горизонтали + 4 по диагонали).")

    h3(doc, "Препятствия")
    body(doc,
        "Узел блокируется, если его координаты попадают в круг препятствия. "
        "Проверка: (x−cx)² + (y−cy)² ≤ r². "
        "Заблокированные узлы не попадают в граф соседей.",
        bold_parts=["(x−cx)² + (y−cy)² ≤ r²"])

    h2(doc, "6.3 Функция стоимости ребра (cost_functions.py)")
    code_block(doc, [
        "# Евклидово расстояние в 3D",
        "def edge_length_3d(a, b):",
        "    dx, dy, dz = b[0]-a[0], b[1]-a[1], b[2]-a[2]",
        "    return math.sqrt(dx*dx + dy*dy + dz*dz)",
        "",
        "# Штраф за уклон (тангенс угла наклона)",
        "def slope_penalty(a, b, alpha):",
        "    horizontal = math.sqrt(dx*dx + dy*dy)",
        "    return alpha * abs(dz / horizontal)  # чем круче — тем дороже",
        "",
        "# Итоговая стоимость ребра",
        "cost = w_len * length + w_slope * slope_penalty + w_risk * risk",
    ], caption="Функции стоимости из cost_functions.py")

    body(doc,
        "dz / horizontal — тангенс угла уклона. При крутом подъёме dz велик, "
        "штраф растёт. Параметр alpha контролирует чувствительность к уклону.",
        bold_parts=["dz / horizontal"])

    add_table(doc,
        headers=["Параметр", "Смысл", "Типичное значение"],
        rows=[
            ["--w-len",   "Вес длины ребра",          "1.0"],
            ["--w-slope", "Вес штрафа за уклон",       "1.0"],
            ["--w-risk",  "Вес «опасной зоны»",        "0.0 (не используется по умолчанию)"],
            ["--alpha",   "Коэффициент уклона",        "1.0"],
        ],
        col_widths=[3.5, 7.0, 5.5],
    )

    h2(doc, "6.4 Алгоритм Dijkstra (search.py)")
    body(doc,
        "Идея: расширяем «волну» от старта, всегда обрабатывая ближайший (самый дешёвый) узел.")

    code_block(doc, [
        "def dijkstra(graph, start, goal, *, weights):",
        "    frontier = [(0.0, start)]     # очередь с приоритетом",
        "    came_from = {}                # откуда пришли в каждый узел",
        "    best_cost = {start: 0.0}     # лучшая стоимость до узла",
        "",
        "    while frontier:",
        "        cost, current = heapq.heappop(frontier)  # самый дешёвый",
        "        if current == goal:",
        "            break               # нашли путь!",
        "        for neighbor in graph.neighbors(current):",
        "            new_cost = cost + graph.edge_cost(current, neighbor, weights)",
        "            if new_cost < best_cost.get(neighbor, inf):",
        "                best_cost[neighbor] = new_cost",
        "                came_from[neighbor] = current",
        "                heapq.heappush(frontier, (new_cost, neighbor))",
    ], caption="Алгоритм Dijkstra — поиск кратчайшего пути")

    body(doc,
        "heapq — минимальная куча (priority queue). При heappop всегда извлекается "
        "элемент с минимальной стоимостью. Сложность: O((V+E) log V).",
        bold_parts=["heapq", "heappop"])

    h3(doc, "Восстановление пути")
    code_block(doc, [
        "# came_from содержит цепочку: goal → ... → start",
        "node = goal",
        "path = [node]",
        "while node != start:",
        "    node = came_from[node]   # идём назад",
        "    path.append(node)",
        "path.reverse()              # разворачиваем: start → goal",
    ], caption="Восстановление пути от финиша к старту")

    h2(doc, "6.5 Алгоритм A*")
    body(doc,
        "A* добавляет эвристику — оценку расстояния от текущего узла до цели. "
        "Вместо «чистой стоимости пути» в очереди хранится: стоимость + эвристика.",
        bold_parts=["эвристику"])
    code_block(doc, [
        "# Эвристика — 3D Евклидово расстояние до финиша",
        "def _heuristic(graph, node, goal):",
        "    return edge_length_3d(graph.point(node), graph.point(goal))",
        "",
        "# Приоритет = реальная стоимость + оценка оставшегося пути",
        "priority = g_cost[neighbor] + _heuristic(graph, neighbor, goal)",
        "heapq.heappush(frontier, (priority, neighbor))",
    ], caption="Ключевое отличие A* от Dijkstra — эвристическая функция")

    body(doc,
        "Эвристика допустима (admissible): она никогда не переоценивает реальное расстояние "
        "(3D Евклид ≤ реальный путь). Это гарантирует нахождение оптимального пути.",
        bold_parts=["допустима (admissible)"])

    h2(doc, "6.6 Сравнение алгоритмов")
    add_table(doc,
        headers=["Алгоритм", "Стратегия", "Посещает узлов", "Гарантия оптимальности"],
        rows=[
            ["Dijkstra", "Расширяется во все стороны равномерно", "Больше", "Да"],
            ["A*",       "Направлен к цели через эвристику",      "Меньше*", "Да"],
        ],
        col_widths=[3.0, 5.5, 3.5, 5.0],
    )
    body(doc,
        "* На сложном рельефе с препятствиями A* может проверять больше узлов, "
        "чем Dijkstra, если эвристика «обманывается» рельефом. Это нормально.")

    h2(doc, "6.7 Визуализация маршрута в Blender")
    body(doc,
        "Маршрут отображается как Curve-объект — трубка, проложенная поверх поверхности. "
        "Зелёная сфера — точка старта, жёлтая сфера — финиш.",
        bold_parts=["Curve-объект", "Зелёная сфера", "жёлтая сфера"])

    code_block(doc, [
        "curve_data.bevel_depth = 0.04    # толщина трубки маршрута",
        "spline.points[idx].co = (x, y, z + 0.04, 1.0)  # +0.04: над поверхностью",
    ], caption="Ключевые параметры Curve-объекта маршрута")

    screenshot_placeholder(doc,
        "Рендер path_paraboloid_astar.png: синяя поверхность параболоида, "
        "красная трубка маршрута от угла до угла, зелёная сфера старта, жёлтая — финиша.")
    screenshot_placeholder(doc,
        "Рендер path_wave_obstacle_astar.png: волновая поверхность, "
        "маршрут огибает круговое препятствие.")

    h2(doc, "6.8 Запуск pathfinding")
    code_block(doc, [
        "# Простой маршрут на параболоиде (A*)",
        "blender --background --python scripts/pathfinding/visualize_path_in_blender.py -- \\",
        "    --function paraboloid --algorithm astar \\",
        "    --start-x -4 --start-y -4 --goal-x 4 --goal-y 4 \\",
        "    --output assets/renders/path_paraboloid_astar.png",
        "",
        "# Маршрут с препятствием (Dijkstra)",
        "blender --background --python scripts/pathfinding/visualize_path_in_blender.py -- \\",
        "    --function wave --algorithm dijkstra \\",
        "    --obstacle-circle 0,0,1.5 \\",
        "    --start-x -4 --start-y -4 --goal-x 4 --goal-y 4",
    ], caption="Команды запуска pathfinding-скрипта")

    doc.add_page_break()


# ────────────────────────────────────────────────────────────
# Глава 7 — Эксперименты
# ────────────────────────────────────────────────────────────

def build_ch7(doc: Document) -> None:
    h1(doc, "Глава 7. Эксперименты и таблицы результатов")

    h2(doc, "7.1 Серия 1: Влияние амплитуды")
    body(doc,
        "Гипотеза: при увеличении A форма становится «острее», стоимость маршрута растёт.")
    add_table(doc,
        headers=["Amplitude (A)", "Функция", "Узлов в пути", "Стоимость", "Наблюдение"],
        rows=[
            ["0.2", "wave, k=1", "___", "___", "Почти плоско, путь прямой"],
            ["0.5", "wave, k=1", "___", "___", "Лёгкий рельеф"],
            ["1.0", "wave, k=1", "___", "___", "Заметные волны"],
            ["2.0", "wave, k=1", "___", "___", "Сильный рельеф"],
            ["3.0", "wave, k=1", "___", "___", "Очень крутые склоны"],
        ],
        col_widths=[3.5, 3.5, 3.0, 3.0, 5.0],
    )
    body_plain(doc,
        "Команда: blender --background --python scripts/pathfinding/visualize_path_in_blender.py "
        "-- --function wave --amplitude X --algorithm astar "
        "--start-x -4 --start-y -4 --goal-x 4 --goal-y 4")

    h2(doc, "7.2 Серия 2: Влияние частоты")
    body(doc,
        "Гипотеза: при увеличении k волн становится больше, путь становится «зигзагообразнее».")
    add_table(doc,
        headers=["Frequency (k)", "Функция", "Узлов в пути", "Стоимость", "Наблюдение"],
        rows=[
            ["0.5", "wave, A=1", "___", "___", "Мало волн, пологие"],
            ["1.0", "wave, A=1", "___", "___", "Базовый случай"],
            ["2.0", "wave, A=1", "___", "___", "Двойная частота"],
            ["3.0", "wave, A=1", "___", "___", "Очень частый рельеф"],
        ],
        col_widths=[3.5, 3.5, 3.0, 3.0, 5.0],
    )

    h2(doc, "7.3 Серия 3: A* против Dijkstra")
    body(doc, "Сравнение алгоритмов на одинаковых задачах:")
    add_table(doc,
        headers=["Алгоритм", "Функция", "Посещено узлов", "Стоимость", "Кто быстрее?"],
        rows=[
            ["Dijkstra", "paraboloid, res=80", "___", "___", "___"],
            ["A*",       "paraboloid, res=80", "___", "___", "___"],
            ["Dijkstra", "wave, res=80",       "___", "___", "___"],
            ["A*",       "wave, res=80",       "___", "___", "___"],
        ],
        col_widths=[3.0, 4.0, 3.5, 3.5, 4.0],
    )
    body_plain(doc, "Данные берутся из строки вывода: [OK] Маршрут найден: N узлов, "
                    "стоимость=X.XXXX, посещено=Y, проверено рёбер=Z.")

    h2(doc, "7.4 Серия 4: Влияние веса уклона")
    body(doc, "Проверяем, как изменяется маршрут при разных w-slope:")
    add_table(doc,
        headers=["--w-slope", "Характер маршрута", "Стоимость", "Наблюдение"],
        rows=[
            ["0",   "Только длина пути (уклон игнорируется)", "___", "___"],
            ["0.5", "Слабый учёт уклона",                     "___", "___"],
            ["1.0", "Стандартный баланс",                     "___", "___"],
            ["3.0", "Сильный штраф за уклон",                 "___", "___"],
            ["5.0", "Максимальное избегание склонов",          "___", "___"],
        ],
        col_widths=[2.5, 6.0, 3.0, 6.5],
    )

    note_box(doc,
        "Заполни таблицы самостоятельно, запуская скрипты с нужными параметрами. "
        "Значения в столбцах «Узлов в пути» и «Стоимость» берутся из вывода терминала.",
        prefix="✏️ Задание")

    doc.add_page_break()


# ────────────────────────────────────────────────────────────
# Глава 8 — Задания
# ────────────────────────────────────────────────────────────

def build_ch8(doc: Document) -> None:
    h1(doc, "Глава 8. Задания для самостоятельной работы")

    h2(doc, "Уровень 1 — Базовый (обязательно)")

    h3(doc, "Задание 1.1 — Пять функций")
    body(doc,
        "Запусти visualize_function.py со всеми пятью функциями. "
        "Сделай скриншот каждой поверхности.")
    body(doc, "Заполни таблицу:")
    add_table(doc,
        headers=["Функция", "На что похоже", "Где встречается в природе/технике"],
        rows=[
            ["paraboloid", "___", "___"],
            ["saddle",     "___", "___"],
            ["wave",       "___", "___"],
            ["ripple",     "___", "___"],
            ["gaussian",   "___", "___"],
        ],
        col_widths=[3.5, 5.0, 8.5],
    )

    h3(doc, "Задание 1.2 — Влияние амплитуды")
    body(doc,
        "Измени --amplitude от 0.5 до 3.0 с шагом 0.5 для функции wave. "
        "Сделай 5 рендеров. Для каждого запиши: «При A=X форма ...».")

    h3(doc, "Задание 1.3 — Сравнение алгоритмов")
    body(doc,
        "Запусти pathfinding на paraboloid с --algorithm astar и --algorithm dijkstra. "
        "Сравни вывод: сколько узлов посетил каждый алгоритм? Почему разница?")

    h2(doc, "Уровень 2 — Средний")

    h3(doc, "Задание 2.1 — Изменение формулы custom")
    body(doc,
        "Функция custom определена в function_library.py:")
    code_block(doc, [
        "def _custom(x, y, config):",
        "    k = config.frequency",
        "    return config.amplitude * (math.sin(k * x) + math.cos(k * y))",
    ], caption="Исходная функция _custom")
    body(doc,
        "Измени формулу — например, добавь: * math.sin(k * (x + y)). "
        "Сделай рендер до и после. Опиши изменение формы.")

    h3(doc, "Задание 2.2 — Добавить новую функцию")
    body(doc,
        "Добавь в FUNCTION_REGISTRY новую функцию. Например, «вулкан»:")
    code_block(doc, [
        "import math",
        "",
        "def _volcano(x, y, config):",
        "    r = math.sqrt(x**2 + y**2)",
        "    return config.amplitude * math.exp(-r) * math.sin(config.frequency * r)",
        "",
        "FUNCTION_REGISTRY['volcano'] = _volcano",
    ], caption="Пример добавления новой функции volcano")
    body(doc,
        "После этого запусти: --function volcano. Получи рендер. "
        "Объясни: почему форма напоминает вулкан?")

    h3(doc, "Задание 2.3 — Заблокированные зоны")
    body(doc,
        "Запусти pathfinding с --blocked-z-gt 0.5 на gaussian. "
        "Объясни: что физически означает этот параметр? "
        "Почему маршрут обходит центр холма?")

    h2(doc, "Уровень 3 — Продвинутый")

    h3(doc, "Задание 3.1 — Новая функция стоимости")
    body(doc,
        "В cost_functions.py стоимость ребра = w_len*длина + w_slope*уклон. "
        "Реализуй штраф только за нисхождение (опасно спускаться, но не подниматься). "
        "Модифицируй slope_penalty в cost_functions.py.")

    h3(doc, "Задание 3.2 — Сравнительный рендер двух алгоритмов")
    body(doc,
        "Напиши скрипт, который запускает A* и Dijkstra на одной поверхности "
        "и отрисовывает оба маршрута в одной Blender-сцене разными цветами "
        "(например, красный для A* и синий для Dijkstra).")

    h3(doc, "Задание 3.3 — Прямоугольное препятствие")
    body(doc,
        "В terrain_graph.py препятствия — только круги. "
        "Добавь поддержку прямоугольных препятствий: задаются как (x_min, x_max, y_min, y_max). "
        "Реализуй _inside_rect_obstacle и добавь параметр --obstacle-rect в CLI.")

    doc.add_page_break()


# ────────────────────────────────────────────────────────────
# Глава 9 — Оформление
# ────────────────────────────────────────────────────────────

def build_ch9(doc: Document) -> None:
    h1(doc, "Глава 9. Оформление результата")

    h2(doc, "9.1 Структура отчёта")
    numbered(doc, "Введение — зачем нужна 3D-визуализация математики (1–2 абзаца)")
    numbered(doc, "Математическая часть — какие функции исследовал, формулы с объяснением")
    numbered(doc, "Технологическая часть — как работает код (ключевые фрагменты с пояснением)")
    numbered(doc, "Экспериментальная часть — таблицы, скриншоты, выводы")
    numbered(doc, "Задача поиска пути — постановка, алгоритм, рендер маршрута, сравнение A* и Dijkstra")
    numbered(doc, "Заключение — что узнал, что было трудным, что можно развить дальше")

    h2(doc, "9.2 Структура презентации")
    add_table(doc,
        headers=["Слайд", "Содержание"],
        rows=[
            ["1",  "Название, ФИО, класс, школа"],
            ["2",  "Мотивация: зачем нужна 3D-математика (фото примеров из жизни)"],
            ["3",  "Что такое z = f(x, y) — объяснение с простым примером"],
            ["4",  "Пять функций — рендеры + формулы в одной таблице"],
            ["5",  "Как работает код — схема: формула → Python → Blender → PNG"],
            ["6",  "Эксперимент: амплитуда — таблица + 3 рендера рядом"],
            ["7",  "Эксперимент: частота — таблица + рендеры"],
            ["8",  "Задача поиска пути — постановка (картинка дрона над рельефом)"],
            ["9",  "A* vs Dijkstra — сравнение + рендер маршрута"],
            ["10", "Маршрут с препятствием — рендер + объяснение обхода"],
            ["11", "Выводы: что узнал, что можно развить"],
            ["12", "Спасибо. Вопросы?"],
        ],
        col_widths=[1.5, 14.5],
    )

    h2(doc, "9.3 Критерии оценки")
    add_table(doc,
        headers=["Критерий", "Описание", "Вес"],
        rows=[
            ["Математика",    "Корректное объяснение z=f(x,y) и функций",        "20%"],
            ["Код",           "Скрипты запускаются, рендеры получены",            "25%"],
            ["Эксперименты",  "Заполненные таблицы с выводами",                   "20%"],
            ["Алгоритм",      "Объяснение поиска пути своими словами",            "20%"],
            ["Оформление",    "Отчёт / презентация выполнены по структуре",       "15%"],
        ],
        col_widths=[4.0, 8.0, 2.0],
    )

    doc.add_page_break()


# ────────────────────────────────────────────────────────────
# Глава 10 — Лабиринт: единый движок поиска на 2D
# ────────────────────────────────────────────────────────────

def build_ch10(doc: Document) -> None:
    h1(doc, "Глава 10. Лабиринт: один и тот же движок на 2D")

    body(doc,
        "В главе 6 мы искали путь на 3D-поверхности z = f(x, y). Теперь покажем, "
        "что тот же самый движок (TerrainGraph + A* / Dijkstra) работает и на плоском "
        "лабиринте — без единой строчки изменений в самих алгоритмах. Ключевая идея: "
        "у рельефа и у лабиринта общий формат данных — 2D-массив ячеек.",
        bold_parts=["тот же самый движок", "общий формат данных"])

    h2(doc, "10.1 Идея и формат данных")
    body(doc,
        "И рельеф, и лабиринт сводятся к матрице H[rows][cols]. В рельефе в ячейках "
        "лежит высота z; в лабиринте — только два значения:")
    bullet(doc, "0 — свободная ячейка (проход)")
    bullet(doc, "1 — стена")
    body(doc,
        "Функция maze_to_terrain_graph превращает эту матрицу в плоский TerrainGraph "
        "(z = 0 во всех точках) и помечает стены как blocked=True. Дальше — обычный "
        "вызов a_star или dijkstra из модуля pathfinding.search.")

    note_box(doc,
        "Важный педагогический момент: студенты видят, что алгоритмы поиска пути "
        "НЕ знают ни про рельеф, ни про лабиринт. Они работают с абстрактным графом. "
        "Меняется только способ построения графа — сама логика A* и Dijkstra остаётся.",
        prefix="📌 Для преподавателя")

    h2(doc, "10.2 Генерация лабиринта (DFS с откатом)")
    body(doc,
        "Используем классический алгоритм «рекурсивного обхода» (recursive backtracking), "
        "реализованный итеративно через стек — чтобы не упираться в лимит рекурсии "
        "Python на больших лабиринтах. Для воспроизводимости принимаем параметр seed: "
        "одинаковый seed → одинаковый лабиринт.")

    code_block(doc, [
        "def generate_maze(rows: int, cols: int, seed: int | None = None) -> MazeGrid:",
        "    # Чётные размеры округляем до нечётных — чтобы корректно легли стены.",
        "    if rows % 2 == 0: rows += 1",
        "    if cols % 2 == 0: cols += 1",
        "",
        "    rng = random.Random(seed)",
        "    maze = [[WALL for _ in range(cols)] for _ in range(rows)]",
        "",
        "    # Стартуем с (1, 1), прорубаем по 2 клетки за шаг.",
        "    stack = [(1, 1)]",
        "    maze[1][1] = FREE",
        "    while stack:",
        "        r, c = stack[-1]",
        "        neighbours = []",
        "        for dr, dc in ((-2, 0), (2, 0), (0, -2), (0, 2)):",
        "            nr, nc = r + dr, c + dc",
        "            if 0 < nr < rows - 1 and 0 < nc < cols - 1 and maze[nr][nc] == WALL:",
        "                neighbours.append((nr, nc, dr, dc))",
        "        if not neighbours:",
        "            stack.pop()",
        "            continue",
        "        nr, nc, dr, dc = rng.choice(neighbours)",
        "        maze[r + dr // 2][c + dc // 2] = FREE  # пробиваем стену между клетками",
        "        maze[nr][nc] = FREE",
        "        stack.append((nr, nc))",
        "    return maze",
    ], caption="scripts/pathfinding/labyrinth.py — generate_maze")

    body(doc,
        "Результат: «идеальный» лабиринт — между любыми двумя клетками существует "
        "ровно один путь, циклов нет. Это ключевое свойство, на котором строится "
        "сравнение A* и Dijkstra в §10.7.")

    h2(doc, "10.3 Перенос лабиринта в TerrainGraph")
    body(doc,
        "TerrainGraph уже умеет работать с заблокированными ячейками. Нам осталось "
        "только заполнить массив высот нулями и пробросить маску стен:")

    code_block(doc, [
        "def maze_to_terrain_graph(maze: MazeGrid, *, cell_size: float = 1.0,",
        "                          connectivity: int = 4) -> TerrainGraph:",
        "    rows = len(maze)",
        "    cols = len(maze[0])",
        "    heights = [[0.0] * cols for _ in range(rows)]",
        "    blocked = [[cell == WALL for cell in row] for row in maze]",
        "    return TerrainGraph(",
        "        heights=heights,",
        "        cell_size=cell_size,",
        "        connectivity=connectivity,  # 4 — без диагоналей, как принято в лабиринтах",
        "        blocked=blocked,",
        "    )",
    ], caption="scripts/pathfinding/labyrinth.py — maze_to_terrain_graph")

    h2(doc, "10.4 Поиск пути")
    body(doc,
        "Тонкая обёртка — строит граф и вызывает нужный алгоритм. Возвращает штатный "
        "SearchResult из pathfinding.search, то есть path, visited_nodes, cost — "
        "абсолютно так же, как и для рельефа.")

    code_block(doc, [
        "def find_path_in_maze(maze, start, goal, *, algorithm: str = \"astar\",",
        "                      cell_size: float = 1.0, connectivity: int = 4,",
        "                      weights: Weights | None = None) -> SearchResult:",
        "    graph = maze_to_terrain_graph(maze, cell_size=cell_size,",
        "                                  connectivity=connectivity)",
        "    weights = weights or Weights(w_length=1.0, w_slope=0.0, w_risk=0.0)",
        "    if algorithm == \"astar\":",
        "        return a_star(graph, start, goal, weights=weights)",
        "    if algorithm == \"dijkstra\":",
        "        return dijkstra(graph, start, goal, weights=weights)",
        "    raise ValueError(f\"Неизвестный алгоритм: {algorithm}\")",
    ], caption="scripts/pathfinding/labyrinth.py — find_path_in_maze")

    h2(doc, "10.5 ASCII-визуализация")
    body(doc,
        "Перед тем, как строить сцену в Blender, удобно посмотреть результат прямо "
        "в терминале. Функция print_maze рендерит лабиринт и путь через Unicode-символы:")

    code_block(doc, [
        "import sys; sys.path.insert(0, \"scripts\")",
        "from pathfinding.labyrinth import (",
        "    generate_maze, find_path_in_maze, maze_start_goal, print_maze,",
        ")",
        "",
        "maze = generate_maze(rows=21, cols=21, seed=42)",
        "start, goal = maze_start_goal(maze)",
        "result = find_path_in_maze(maze, start, goal, algorithm=\"astar\")",
        "print(f\"Путь: {len(result.path)} ячеек, посещено {result.visited_nodes}\")",
        "print_maze(maze, path=result.path, start=start, goal=goal)",
    ], caption="Минимальный запуск без Blender")

    body(doc,
        "Повторный запуск с тем же seed=42 даст ровно тот же лабиринт и тот же путь — "
        "это требование воспроизводимости из ТЗ проекта (§12.5 AI_AGENT_HANDOFF).")

    h2(doc, "10.6 Сцена в Blender")
    body(doc,
        "Скрипт scripts/pathfinding/visualize_labyrinth_in_blender.py собирает сцену "
        "из четырёх типов объектов:")
    numbered(doc, "Плоский пол — Plane размером rows×cols")
    numbered(doc, "Стены — один объединённый mesh со всеми кубиками (быстро, меньше объектов в сцене)")
    numbered(doc, "Маршрут — Curve с эмиссионным материалом (светится)")
    numbered(doc, "Маркеры START/GOAL — сферы зелёного и жёлтого цвета")

    body(doc,
        "Ключевой фрагмент — построение меша стен. Вместо сотни отдельных кубиков "
        "мы собираем все вершины и грани в один массив и вызываем mesh.from_pydata() "
        "один раз. На лабиринте 41×41 это быстрее в десятки раз:")

    code_block(doc, [
        "def build_walls_mesh(maze, *, cell_size: float, wall_height: float):",
        "    verts, faces = [], []",
        "    for r, row in enumerate(maze):",
        "        for c, cell in enumerate(row):",
        "            if cell != WALL:",
        "                continue",
        "            x = (c - len(row) / 2) * cell_size",
        "            y = (r - len(maze) / 2) * cell_size",
        "            base = len(verts)",
        "            # 8 вершин кубика",
        "            for dx in (0, cell_size):",
        "                for dy in (0, cell_size):",
        "                    for dz in (0, wall_height):",
        "                        verts.append((x + dx, y + dy, dz))",
        "            # 6 граней (по 4 вершины)",
        "            faces.extend([",
        "                (base+0, base+1, base+3, base+2),  # нижняя",
        "                (base+4, base+6, base+7, base+5),  # верхняя",
        "                (base+0, base+4, base+5, base+1),  # передняя",
        "                (base+2, base+3, base+7, base+6),  # задняя",
        "                (base+0, base+2, base+6, base+4),  # левая",
        "                (base+1, base+5, base+7, base+3),  # правая",
        "            ])",
        "    mesh = bpy.data.meshes.new(\"MazeWalls\")",
        "    mesh.from_pydata(verts, [], faces)",
        "    mesh.update()",
        "    return bpy.data.objects.new(\"MazeWalls\", mesh)",
    ], caption="scripts/pathfinding/visualize_labyrinth_in_blender.py — build_walls_mesh")

    body(doc, "Запуск в headless-режиме:")
    code_block(doc, [
        "blender --background --python scripts/pathfinding/visualize_labyrinth_in_blender.py -- \\",
        "    --rows 21 --cols 21 --seed 42 --algorithm astar \\",
        "    --output assets/renders/labyrinth_21_astar.png",
    ])

    h2(doc, "10.7 A* vs Dijkstra на идеальном лабиринте")
    body(doc,
        "В идеальном лабиринте между двумя клетками существует ровно один путь, "
        "поэтому длина маршрута у A* и Dijkstra совпадает — этот факт закреплён "
        "тестом test_astar_and_dijkstra_same_length_on_perfect_maze. Разница "
        "проявляется в числе посещённых узлов: эвристика A* направляет поиск "
        "в сторону цели, Dijkstra же разворачивает волну во все стороны.",
        bold_parts=["длина маршрута у A* и Dijkstra совпадает", "числе посещённых узлов"])

    add_table(doc,
        headers=["Размер", "seed", "Длина пути", "A* посетил", "Dijkstra посетил"],
        rows=[
            ["21×21",  "42",  "101", "125", "135"],
            ["25×25",  "7",   "129", "168", "189"],
            ["31×31",  "123", "189", "241", "278"],
        ],
        col_widths=[3.0, 2.0, 3.5, 3.5, 3.5],
    )

    body(doc,
        "Наблюдение: в лабиринте без циклов выигрыш A* скромный (10–20% по visited_nodes), "
        "потому что эвристика часто упирается в стену и приходится откатываться. "
        "Если «пробить» часть стен и добавить циклы — A* начнёт выигрывать сильнее. "
        "Это материал для проекта C из рабочего листа студента (docs/student_exercises.md).")

    h2(doc, "10.8 Тесты")
    body(doc,
        "Модуль покрыт 14 тестами в tests/test_labyrinth.py. Они не требуют Blender "
        "и запускаются одной командой:")
    code_block(doc, ["python -m unittest tests.test_labyrinth -v"])

    add_table(doc,
        headers=["Группа", "Что проверяется"],
        rows=[
            ["Генерация",  "Воспроизводимость seed, чётные→нечётные размеры, границы всегда стены"],
            ["Формат",     "В ячейках только 0/1, start и goal всегда свободны"],
            ["Поиск",      "A* и Dijkstra находят путь, path идёт только по свободным клеткам"],
            ["Смежность",  "Каждый следующий шаг пути — сосед предыдущего (4-связность)"],
            ["Эквивалентность", "На идеальном лабиринте длина пути A* = длина пути Dijkstra"],
            ["Конвертация", "maze_path_to_scene_points даёт координаты с правильным шагом cell_size"],
        ],
        col_widths=[4.0, 12.0],
    )

    h2(doc, "10.9 Что дальше")
    bullet(doc, "Добавить «дорогие» ячейки (не стены, а высокая стоимость) и показать, как w_risk влияет на маршрут.")
    bullet(doc, "Сгенерировать лабиринт с циклами (убрать часть стен после DFS) — и показать, что A* начнёт заметно выигрывать по visited_nodes.")
    bullet(doc, "Анимация пошагового поиска: сохранять visited_nodes на каждом шаге и запекать keyframes в Blender.")
    bullet(doc, "Скрестить лабиринт и рельеф: пустить путь по лабиринту, у которого стены — холмы на поверхности z=f(x,y).")

    doc.add_page_break()


# ────────────────────────────────────────────────────────────
# Приложение А — Типичные ошибки
# ────────────────────────────────────────────────────────────

def build_appendix_a(doc: Document) -> None:
    h1(doc, "Приложение А. Типичные ошибки и их решения")

    add_table(doc,
        headers=["Ошибка", "Причина", "Решение"],
        rows=[
            ["ModuleNotFoundError: No module named 'bpy'",
             "Скрипт запущен в обычном Python, не в Blender",
             "Использовать blender --background --python ..."],
            ["AttributeError: 'Mesh' object has no attribute 'calc_normals'",
             "Метод удалён в Blender 4+",
             "Удалить строку mesh.calc_normals() из кода"],
            ["ValueError: Неизвестная функция 'parabolid'",
             "Опечатка в имени функции",
             "Правильно: paraboloid, saddle, wave, ripple, gaussian, custom"],
            ["Error: Start-узел в блокированной зоне",
             "Точка старта внутри препятствия",
             "Изменить --start-x / --start-y за пределы препятствия"],
            ["ValueError: Маршрут не найден",
             "Путь полностью заблокирован",
             "Уменьшить радиус препятствия или изменить --blocked-z-gt"],
            ["Чёрный рендер",
             "Нет источника света или камеры",
             "Скрипт создаёт их автоматически; запустить ensure_camera_and_light()"],
            ["Поверхность серая (без цвета)",
             "Режим Viewport — Solid вместо Material Preview",
             "Переключить в правом верхнем углу Viewport"],
            ["Маршрут не виден на рендере",
             "Curve-объект не включён в рендер",
             "Properties → Object Properties → включить иконку камеры"],
        ],
        col_widths=[5.0, 5.0, 6.0],
    )

    doc.add_page_break()


# ────────────────────────────────────────────────────────────
# Приложение Б — Глоссарий
# ────────────────────────────────────────────────────────────

def build_appendix_b(doc: Document) -> None:
    h1(doc, "Приложение Б. Глоссарий терминов")

    terms = [
        ("mesh", "Трёхмерный объект, состоящий из вершин, рёбер и граней"),
        ("вершина (vertex)", "Точка в 3D-пространстве с координатами (x, y, z)"),
        ("грань (face)", "Плоский многоугольник, ограниченный вершинами"),
        ("colormap", "Цветовая карта: числовое значение → цвет (например, z → синий/зелёный/красный)"),
        ("Geometry Nodes", "Система процедурной параметрической геометрии в Blender"),
        ("нод (node)", "Блок вычислений в нод-редакторе Blender (Shader Editor, GN Editor)"),
        ("граф", "Математическая структура: множество узлов и рёбер с весами"),
        ("очередь с приоритетом (heap)", "Структура данных для быстрого извлечения минимального элемента"),
        ("эвристика", "Приближённая оценка расстояния до цели в алгоритме A*"),
        ("уклон (slope)", "Отношение перепада высоты к горизонтальному расстоянию"),
        ("bpy", "Python API Blender — библиотека для управления Blender из Python"),
        ("dataclass", "Декоратор Python для автоматического создания класса с полями"),
        ("frozen", "Неизменяемость объекта dataclass (frozen=True)"),
        ("resolution", "Число шагов сетки по каждой оси; resolution=60 → 61×61 вершин"),
        ("amplitude (A)", "Высота/масштаб значений функции"),
        ("frequency (k)", "Частота волн — сколько периодов на единицу длины"),
        ("sigma (σ)", "Ширина гауссова купола; большое σ → широкая пологая горка"),
        ("8-связность", "Граф, где каждый узел соединён с 8 соседями (вкл. диагонали)"),
    ]

    for term, definition in terms:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        r1 = p.add_run(f"{term}  —  ")
        r1.bold = True
        r1.font.size = Pt(11)
        r1.font.color.rgb = COLOR_ACCENT
        r2 = p.add_run(definition)
        r2.font.size = Pt(11)


# ────────────────────────────────────────────────────────────
# Настройки страницы
# ────────────────────────────────────────────────────────────

def setup_page(doc: Document) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    section = doc.sections[0]
    section.page_width  = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.0)
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # Нумерация страниц в нижнем колонтитуле
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.clear()
    run = footer_para.add_run("Страница ")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x99)
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run2 = footer_para.add_run()
    run2.font.size = Pt(9)
    run2._r.append(fldChar1)
    run2._r.append(instrText)
    run2._r.append(fldChar2)

    # Колонтитул сверху
    header = section.header
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = header_para.add_run("STEM-проект: 3D-визуализация математических поверхностей в Blender")
    r.font.size = Pt(8)
    r.italic = True
    r.font.color.rgb = RGBColor(0x88, 0x88, 0x99)


# ────────────────────────────────────────────────────────────
# Основная функция
# ────────────────────────────────────────────────────────────

def build_docx() -> None:
    print("[START] Создание Методичка_подробная.docx ...")

    doc = Document()
    setup_page(doc)

    # Убираем дефолтный стиль Normal (шрифт и межстрочный интервал)
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    from docx.shared import Pt as _Pt
    style.paragraph_format.line_spacing = _Pt(14)

    build_title_page(doc)
    build_toc(doc)
    build_ch0(doc)
    build_ch1(doc)
    build_ch2(doc)
    build_ch3(doc)
    build_ch4(doc)
    build_ch5(doc)
    build_ch6(doc)
    build_ch7(doc)
    build_ch8(doc)
    build_ch9(doc)
    build_ch10(doc)
    build_appendix_a(doc)
    build_appendix_b(doc)

    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc.save(OUTPUT_PATH)
    size_kb = os.path.getsize(OUTPUT_PATH) // 1024
    print(f"[DONE] Сохранено: {OUTPUT_PATH}  ({size_kb} KB)")


if __name__ == "__main__":
    build_docx()
